"""Balanced FOR_CLIENT.docx. Simple words, normal English, all-black text."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor


BLACK = RGBColor(0, 0, 0)


def _force_black(run):
    run.font.color.rgb = BLACK


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        _force_black(r)
    return h


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    _force_black(run)
    return p


def bullet(doc, label, body):
    p = doc.add_paragraph(style='List Bullet')
    b = p.add_run(label + ': ')
    b.bold = True
    b.font.size = Pt(11)
    _force_black(b)
    rest = p.add_run(body)
    rest.font.size = Pt(11)
    _force_black(rest)


def code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    _force_black(run)
    p.paragraph_format.left_indent = Inches(0.25)


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    doc.add_heading('CodeLogic Tech Stack and How to Run', level=0)
    para(doc, 'A simple guide written for people who are not developers.', italic=True)

    # ---------- About the project ----------
    heading(doc, 'About the project')
    para(
        doc,
        'CodeLogic is a website where students can learn programming through '
        'lessons and quizzes. The project is built in two main parts: a '
        'backend that runs on a server, and a frontend that runs in the '
        'browser. Both parts work together to provide the full experience.',
    )

    # ---------- Backend ----------
    heading(doc, 'Backend (the server side)')
    para(
        doc,
        'The backend stores all the data and handles the rules of the system. '
        'For example, it checks logins, saves quiz scores, and gives the '
        'website the data it needs to show.',
    )
    bullet(doc, 'Python',
           'The programming language used to write the backend.')
    bullet(doc, 'Django',
           'The framework that runs the backend. It is similar to Laravel in '
           'PHP, but for Python. It also gives us a ready-made admin panel.')
    bullet(doc, 'Django REST Framework',
           'A small add-on for Django that builds the API. The API is how the '
           'website asks the server for data.')
    bullet(doc, 'PostgreSQL',
           'The database. It saves all users, quiz questions, lessons, scores, '
           'and certificates.')

    # ---------- Frontend ----------
    heading(doc, 'Frontend (the website)')
    para(
        doc,
        'The frontend is what students and admins see in their browser. It '
        'shows the lessons, the quiz screen, the leaderboard, and so on. It '
        'talks to the backend to get and save data.',
    )
    bullet(doc, 'Next.js',
           'The framework used to build the website. It is built on top of '
           'React.')
    bullet(doc, 'React',
           'The library that draws the screens, buttons, and forms.')
    bullet(doc, 'TypeScript',
           'A safer version of JavaScript. It checks for mistakes before the '
           'code runs.')
    bullet(doc, 'Tailwind CSS',
           'Used for the design, colors, and layout of the website.')

    # ---------- Hosting ----------
    heading(doc, 'Where it is hosted')
    para(
        doc,
        'The two parts of the project live on different services online. The '
        'code itself is kept on GitHub.',
    )
    bullet(doc, 'Vercel', 'Hosts the website (the frontend).')
    bullet(doc, 'DigitalOcean', 'Hosts the backend and the database.')
    bullet(doc, 'GitHub', 'Stores all the source code.')

    # ---------- How to run ----------
    doc.add_page_break()
    heading(doc, 'How to run the project on your computer')
    para(
        doc,
        'These steps explain how to run a personal copy of CodeLogic on a '
        'local computer. This is useful for testing changes without '
        'affecting the live website.',
    )

    para(doc, 'Programs to install first:', bold=True)
    bullet(doc, 'Python 3.10 or newer', 'https://www.python.org/downloads/')
    bullet(doc, 'Node.js 18 or newer', 'https://nodejs.org/')
    bullet(doc, 'Git', 'https://git-scm.com/downloads')

    para(doc, 'Step 1. Download the source code', bold=True)
    para(doc, 'Open a terminal and run:')
    code(doc, 'git clone <repository-url>\ncd codelogic')

    para(doc, 'Step 2. Start the backend', bold=True)
    para(doc, 'In the same terminal, run these commands one by one:')
    code(
        doc,
        'cd backend\n'
        'python -m venv venv\n'
        'venv\\Scripts\\activate          (on Windows)\n'
        'source venv/bin/activate       (on Mac or Linux)\n'
        'pip install -r requirements.txt\n'
        'python manage.py migrate\n'
        'python manage.py runserver'
    )
    para(
        doc,
        'When this finishes, the backend is running at '
        'http://localhost:8000. Keep this terminal open.',
        italic=True,
    )

    para(doc, 'Step 3. Start the frontend', bold=True)
    para(
        doc,
        'Open a second terminal window. The first one must stay open. Then '
        'run:',
    )
    code(doc, 'cd frontend\nnpm install\nnpm run dev')
    para(
        doc,
        'The website is now running at http://localhost:3000. Open that '
        'address in any browser to use the app.',
        italic=True,
    )

    para(doc, 'Step 4. Create an admin account (only once)', bold=True)
    para(
        doc,
        'To use the admin panel locally, an admin user must be created. In '
        'a separate terminal, run:',
    )
    code(doc, 'cd backend\npython manage.py createsuperuser')
    para(
        doc,
        'Set an email, username, and password. After that, log in at '
        'http://localhost:8000/admin.',
        italic=True,
    )

    # ---------- Notes ----------
    heading(doc, 'Important notes')
    bullet(
        doc, 'Both terminals must stay open',
        'The backend runs in one terminal and the frontend in another. If '
        'either is closed, that part of the app stops working.',
    )
    bullet(
        doc, 'Nothing else needs to be installed',
        'The pip install and npm install commands above download every '
        'other package the project needs.',
    )
    bullet(
        doc, 'How is it different from a PHP or Laravel project?',
        'The idea is the same. Django plays the same role for Python as '
        'Laravel plays for PHP. The main difference is that the website '
        'is a separate Next.js project instead of being part of the backend.',
    )

    primary = 'FOR_CLIENT.docx'
    try:
        doc.save(primary)
        print(f'Wrote {primary}')
    except PermissionError:
        fallback = 'FOR_CLIENT_v2.docx'
        doc.save(fallback)
        print(f'{primary} is open in Word, wrote {fallback} instead.')


if __name__ == '__main__':
    main()
