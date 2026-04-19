"""
Convert every .md file under docs/ into a .docx sibling.
Handles: headings, paragraphs, bullet/numbered lists, tables,
code blocks, inline code, bold, italic, and links.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
MD_ROOT = ROOT / "md"
DOCX_ROOT = ROOT / "docx"

INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
ITALIC = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_runs(paragraph, text):
    """Parse inline markdown in text and append runs to paragraph."""
    # First: pull links out so their text doesn't get reformatted
    # Work with a token list
    tokens = [("text", text)]

    def split_tokens(pattern, kind):
        new_tokens = []
        for t_type, t_val in tokens:
            if t_type != "text":
                new_tokens.append((t_type, t_val))
                continue
            idx = 0
            for m in pattern.finditer(t_val):
                if m.start() > idx:
                    new_tokens.append(("text", t_val[idx:m.start()]))
                if kind == "link":
                    new_tokens.append(("link", (m.group(1), m.group(2))))
                else:
                    new_tokens.append((kind, m.group(1)))
                idx = m.end()
            if idx < len(t_val):
                new_tokens.append(("text", t_val[idx:]))
        return new_tokens

    tokens = split_tokens(LINK, "link")
    tokens = split_tokens(INLINE_CODE, "code")
    tokens = split_tokens(BOLD, "bold")
    tokens = split_tokens(ITALIC, "italic")

    for t_type, t_val in tokens:
        if t_type == "text":
            paragraph.add_run(t_val)
        elif t_type == "bold":
            r = paragraph.add_run(t_val)
            r.bold = True
        elif t_type == "italic":
            r = paragraph.add_run(t_val)
            r.italic = True
        elif t_type == "code":
            r = paragraph.add_run(t_val)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif t_type == "link":
            text_val, _url = t_val
            r = paragraph.add_run(text_val)
            r.font.color.rgb = RGBColor(0x06, 0x5F, 0xD0)
            r.underline = True


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("\n".join(code_lines))
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Light gray shading on the paragraph
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F3F3")
    p_pr.append(shd)


def add_table(doc, header, rows):
    cols = max(len(header), max((len(r) for r in rows), default=0))
    if cols == 0:
        return
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Light Grid Accent 1"

    # Header
    for i in range(cols):
        cell = table.rows[0].cells[i]
        text = header[i] if i < len(header) else ""
        p = cell.paragraphs[0]
        add_runs(p, text)
        for run in p.runs:
            run.bold = True
        shade_cell(cell, "4A4A7A")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Body
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            add_runs(p, text)


def convert(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    print(f"  {md_path.relative_to(MD_ROOT)} -> docx/{docx_path.relative_to(DOCX_ROOT)}")
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pbdr.append(bottom)
            p_pr.append(pbdr)
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?[\s\-:|]+\|[\s\-:|]*$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, header, rows)
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_heading(level=min(level, 4))
            add_runs(p, text)
            i += 1
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_runs(p, text)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph (may wrap over multiple lines until blank)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            re.match(r"^(#{1,6})\s", lines[i])
            or lines[i].strip().startswith("```")
            or re.match(r"^\s*[-*]\s+", lines[i])
            or re.match(r"^\s*\d+\.\s+", lines[i])
            or "|" in lines[i]
            or re.match(r"^---+\s*$", lines[i])
        ):
            para_lines.append(lines[i])
            i += 1

        p = doc.add_paragraph()
        add_runs(p, " ".join(l.strip() for l in para_lines))

    doc.save(str(docx_path))


def main():
    md_files = sorted(MD_ROOT.rglob("*.md"))
    print(f"Converting {len(md_files)} markdown file(s) to .docx...")
    skipped = []
    for md in md_files:
        rel = md.relative_to(MD_ROOT)
        docx = (DOCX_ROOT / rel).with_suffix(".docx")
        docx.parent.mkdir(parents=True, exist_ok=True)
        try:
            convert(md, docx)
        except PermissionError:
            skipped.append(str(rel))
            print(f"  SKIPPED (file in use — close it in Word): {rel}")
    print("Done.")
    if skipped:
        print(f"\n{len(skipped)} file(s) skipped. Close them in Word and rerun.")


if __name__ == "__main__":
    main()
