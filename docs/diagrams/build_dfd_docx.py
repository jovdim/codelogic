"""Build CodeLogic_DFD.docx - full DFD set (Level 0, Level 1, Level 2 x6)."""
from pathlib import Path
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT, WD_SECTION

HERE = Path(__file__).parent
OUT = HERE.parent / "CodeLogic_DFD_v3.docx"

doc = Document()

# ---------------------------------------------------------------------------
# First section: landscape, for the wide Level 0 context diagram.
# ---------------------------------------------------------------------------
s = doc.sections[0]
nw, nh = s.page_height, s.page_width
s.orientation = WD_ORIENT.LANDSCAPE
s.page_width = nw
s.page_height = nh
s.left_margin = Inches(0.5)
s.right_margin = Inches(0.5)
s.top_margin = Inches(0.5)
s.bottom_margin = Inches(0.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def _heading(text: str, size_pt: int = 18, keep_with_next: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if keep_with_next:
        # Glue heading to the image paragraph below so Word can't strand
        # the heading on an otherwise-empty page when the image is tall.
        p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size_pt)


def _fit_image(image_path: Path, max_w_in: float, max_h_in: float):
    """Return (width_in, height_in) that fits the image within the given
    page area while preserving aspect ratio. Width is the primary limit;
    height kicks in only when the image is taller than the box allows."""
    with PILImage.open(image_path) as im:
        w_px, h_px = im.size
    aspect = h_px / w_px
    # First try width-constrained.
    width_in = max_w_in
    height_in = width_in * aspect
    if height_in > max_h_in:
        # Image is too tall for the box - constrain by height instead.
        height_in = max_h_in
        width_in = height_in / aspect
    return width_in, height_in


def _figure(image_path: Path, fig_num: int, caption: str,
            max_w_in: float, max_h_in: float):
    """Insert an image figure that always fits the available page area.

    max_w_in / max_h_in describe the usable box on the page AFTER the
    headings and caption have taken their share - so the image cannot
    overflow onto a blank next page.
    """
    width_in, height_in = _fit_image(image_path, max_w_in, max_h_in)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Glue the image to the caption that follows so they always stay on
    # the same page.
    p_img.paragraph_format.keep_with_next = True
    r = p_img.add_run()
    r.add_picture(str(image_path), width=Inches(width_in),
                  height=Inches(height_in))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = cap.add_run(f"Figure {fig_num}. ")
    fr.bold = True
    fr.italic = True
    fr.font.size = Pt(11)
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(11)


def _portrait_section():
    """Append a new portrait section so subsequent figures get the
    taller page format the Level 1 / Level 2 diagrams need."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    nw2, nh2 = sec.page_height, sec.page_width
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = nh2
    sec.page_height = nw2
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)


def _page_break():
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


# Page-area budgets. A US Letter page minus 0.5" margins leaves 7.5" x 10"
# of usable space. Headings + caption together take ~1.4", so the image
# itself gets the remainder. Landscape pages flip the dimensions.
PORTRAIT_MAX_W = 7.5
PORTRAIT_MAX_H = 8.6     # 10.0 usable - ~1.4 for two headings + caption
LANDSCAPE_MAX_W = 10.0
LANDSCAPE_MAX_H = 6.1    # 7.5 usable - ~1.4 for heading + caption

# ---------------------------------------------------------------------------
# Figure 1: Level 0 (context diagram) - landscape
# ---------------------------------------------------------------------------
_heading("CodeLogic: Data Flow Diagram - Level 0", keep_with_next=True)
_figure(
    HERE / "03_dfd_level0.png",
    1,
    "Level 0 Data Flow Diagram (Context Diagram) of the CodeLogic "
    "Quiz Platform showing the four external entities (Guest, Student, "
    "Admin, Email Service, PDF Service) and the high-level data flows "
    "into and out of the system.",
    max_w_in=LANDSCAPE_MAX_W, max_h_in=LANDSCAPE_MAX_H,
)

# ---------------------------------------------------------------------------
# Figure 2: Level 1 (system decomposition) - portrait
# ---------------------------------------------------------------------------
_portrait_section()
_heading("CodeLogic: Data Flow Diagram - Level 1", keep_with_next=True)
_figure(
    HERE / "04_dfd_level1.png",
    2,
    "Level 1 Data Flow Diagram showing the six top-level processes of "
    "the CodeLogic platform (User & Auth Management, Content Catalog, "
    "Quiz Engine, Progress & Gamification, Certificate Management, and "
    "Learning Resources) and their primary data stores.",
    max_w_in=PORTRAIT_MAX_W, max_h_in=PORTRAIT_MAX_H,
)

# ---------------------------------------------------------------------------
# Figures 3-8: Level 2 decompositions, one per Level 1 process.
# Each lives on its own page so a reader can study the sub-process layout
# without scrolling between two figures at once.
# ---------------------------------------------------------------------------
LEVEL2_FIGURES = [
    (
        "04_dfd_level2_1_auth.png",
        "Process 1.0 - User & Auth Management",
        "Level 2 decomposition of Process 1.0 (User & Auth Management). "
        "Shows the nine sub-processes that handle registration, email "
        "verification, login with lockout, account unlock, password reset, "
        "logout, profile / avatar / password updates, and the post-login "
        "face-snapshot capture.",
    ),
    (
        "04_dfd_level2_2_content.png",
        "Process 2.0 - Content Catalog",
        "Level 2 decomposition of Process 2.0 (Content Catalog). Shows "
        "how guests and students browse categories, category details, and "
        "topic details (with per-user progress), and how admins manage "
        "categories, topics, questions, and lessons through the Django "
        "admin interface.",
    ),
    (
        "04_dfd_level2_3_quiz.png",
        "Process 3.0 - Quiz Engine",
        "Level 2 decomposition of Process 3.0 (Quiz Engine). Shows the "
        "four sub-processes of the gameplay loop: starting an attempt "
        "(serving questions + lessons), validating and scoring each "
        "answer, handling the 30-second per-question timeout, and "
        "completing the quiz (final scoring + stars + XP).",
    ),
    (
        "04_dfd_level2_4_progress.png",
        "Process 4.0 - Progress & Gamification",
        "Level 2 decomposition of Process 4.0 (Progress & Gamification). "
        "Shows how hearts regenerate over time, how XP is awarded with "
        "bonuses, how the daily streak is recomputed, how per-topic "
        "progress is bumped, and the three read-only dashboards (user "
        "stats, daily stats / challenges, and leaderboard).",
    ),
    (
        "04_dfd_level2_5_certificates.png",
        "Process 5.0 - Certificate Management",
        "Level 2 decomposition of Process 5.0 (Certificate Management). "
        "Shows how a student's earned certificates are listed, how the "
        "certificate HTML is assembled, how it is rendered to PDF by "
        "either headless Chrome or WeasyPrint, and how admins can preview "
        "any user's certificate from the admin panel.",
    ),
    (
        "04_dfd_level2_6_resources.png",
        "Process 6.0 - Learning Resources",
        "Level 2 decomposition of Process 6.0 (Learning Resources). Shows "
        "how guests and students browse, search, and filter the resource "
        "library, how a resource detail view increments its view counter, "
        "and how admins upload, edit, and remove resources.",
    ),
]

for i, (filename, parent_label, caption) in enumerate(LEVEL2_FIGURES, start=3):
    _page_break()
    _heading("CodeLogic: Data Flow Diagram - Level 2", size_pt=18,
             keep_with_next=True)
    _heading(parent_label, size_pt=13, keep_with_next=True)
    _figure(HERE / filename, i, caption,
            max_w_in=PORTRAIT_MAX_W, max_h_in=PORTRAIT_MAX_H)


doc.save(OUT)
print(f"Wrote {OUT}")
