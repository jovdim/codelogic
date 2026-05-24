"""Build CodeLogic_Diagrams.docx - flowchart, use case, DFD L0, DFD L1."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
OUT = HERE.parent / "CodeLogic_Diagrams_v5.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CodeLogic: System Diagrams")
run.bold = True
run.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Flowchart, Use Case Diagram, and Data Flow Diagrams")
sub_run.italic = True
sub_run.font.size = Pt(12)
sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()


def add_figure(image_path: Path, figure_num: int, caption_text: str, max_width_in: float = 6.5):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Inches(max_width_in))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_run = cap.add_run(f"Figure {figure_num} ")
    fig_run.bold = True
    fig_run.italic = True
    fig_run.font.size = Pt(11)
    cap_run = cap.add_run(caption_text)
    cap_run.italic = True
    cap_run.font.size = Pt(11)

    doc.add_paragraph()


add_figure(
    HERE / "01_flowchart_student.png",
    1,
    "System Flowchart of the Student Side of CodeLogic",
)

doc.add_page_break()

add_figure(
    HERE / "02_use_case.png",
    2,
    "Use Case Diagram of the CodeLogic Quiz Platform",
)

doc.add_page_break()

add_figure(
    HERE / "03_dfd_level0.png",
    3,
    "Level 0 Data Flow Diagram (Context Diagram) of the CodeLogic Quiz Platform",
)

doc.add_page_break()

add_figure(
    HERE / "04_dfd_level1.png",
    4,
    "Level 1 Data Flow Diagram: User, Content, Quiz, Progress, Certificate, and Resource Management",
)

doc.save(OUT)
print(f"Wrote {OUT}")
