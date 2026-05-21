"""Build CodeLogic_ERD.docx — ERD only, portrait."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
OUT = HERE.parent / "CodeLogic_ERD_v3.docx"

doc = Document()

s = doc.sections[0]
s.left_margin = Inches(0.5)
s.right_margin = Inches(0.5)
s.top_margin = Inches(0.5)
s.bottom_margin = Inches(0.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CodeLogic: Entity Relationship Diagram")
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()

p_img = doc.add_paragraph()
p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_img.add_run()
r.add_picture(str(HERE / "05_erd.png"), width=Inches(7.5))

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = cap.add_run("Figure 1 ")
fr.bold = True
fr.italic = True
fr.font.size = Pt(11)
cr = cap.add_run("Entity Relationship Diagram of the CodeLogic Quiz Platform")
cr.italic = True
cr.font.size = Pt(11)

doc.save(OUT)
print(f"Wrote {OUT}")
