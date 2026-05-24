"""Build CodeLogic_Flowchart.docx - flowchart only."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

HERE = Path(__file__).parent
OUT = HERE.parent / "CodeLogic_Flowchart_v3.docx"

doc = Document()

section = doc.sections[0]
new_w, new_h = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_w
section.page_height = new_h
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CodeLogic: System Flowchart")
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph()

p_img = doc.add_paragraph()
p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_img.add_run()
run.add_picture(str(HERE / "01_flowchart_student.png"), width=Inches(10.0))

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_run = cap.add_run("Figure 1 ")
fig_run.bold = True
fig_run.italic = True
fig_run.font.size = Pt(11)
cap_run = cap.add_run("System Flowchart of the CodeLogic Quiz Platform")
cap_run.italic = True
cap_run.font.size = Pt(11)

doc.save(OUT)
print(f"Wrote {OUT}")
